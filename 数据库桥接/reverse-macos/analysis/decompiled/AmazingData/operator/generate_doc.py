# Source Generated with Decompyle++
# File: generate_doc.pyc (Python 3.12)

'''
算子函数说明文档生成器
生成Word格式的算子函数说明文档，包含四类函数：
1. 数学函数 (MathFunction)
2. 统计函数 (StatisticsFunction)
3. 时序函数 (TimeSeriesFunction)
4. 截面函数 (CrossSectionFunction)
'''
import inspect
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from math_function import MathFunction
from statistics_function import StatisticsFunction
from time_series_function import TimeSeriesFunction
from cross_section_function import CrossSectionFunction

def parse_docstring(docstring):
    '''解析函数的docstring，提取描述和用法'''
    if not docstring:
        return ('', '')
    lines = docstring.strip().split('\n')
    description = ''
    usage = ''
    for line in lines:
        line = line.strip()
        if line.startswith('用法:') or line.startswith('用法：'):
            usage = line
            continue
        if line.startswith(':param'):
            continue
        if line.startswith(':return'):
            continue
        if not line:
            continue
        if description:
            continue
        description = line
    return (description, usage)


def get_function_signature(func):
    '''获取函数签名'''
    
    try:
        sig = inspect.signature(func)
        params = []
        for name, param in sig.parameters.items():
            if param.annotation != inspect.Parameter.empty:
                annotation = param.annotation
                params.append(f'''{name}: {type_name}''')
                continue
            params.append(name)
        return f'''({', '.join(params)})'''
    except:
        return '()'



def get_functions_from_class(cls):
    '''从类中获取所有公开的静态方法'''
    functions = []
    for name, method in inspect.getmembers(cls, predicate = inspect.isfunction):
        if name.startswith('_'):
            continue
        functions.append((name, method))
    return sorted(functions, key = (lambda x: x[0]))


def set_run_font(run, font_name, font_size, bold = ('仿宋', 10.5, False)):
    '''设置run的字体'''
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold


def create_document():
    '''创建Word文档'''
    doc = Document()
    doc.styles['Normal'].font.name = '仿宋'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    doc.styles['Normal'].font.size = Pt(10.5)
    return doc


def add_title(doc, text, level = (1,)):
    '''添加标题'''
    if level == 0:
        heading = doc.add_heading(text, level = 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            set_run_font(run, '黑体', 22, True)
        return heading
    if None == 1:
        heading = doc.add_heading(text, level = 1)
        for run in heading.runs:
            set_run_font(run, '黑体', 16, True)
        return heading
    if None == 2:
        heading = doc.add_heading(text, level = 2)
        for run in heading.runs:
            set_run_font(run, '黑体', 14, True)
        return heading
    heading = None.add_heading(text, level = 3)
    for run in heading.runs:
        set_run_font(run, '仿宋', 12, True)
    return heading


def add_function_table(doc, functions):
    '''添加函数列表表格'''
    table = doc.add_table(rows = 1, cols = 3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    headers = [
        '序号',
        '函数名称',
        '函数用法']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, '仿宋', 10.5, True)
    for name, method in enumerate(functions, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        for paragraph in row_cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = name
        for paragraph in row_cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        (description, usage) = parse_docstring(method.__doc__)
        if usage:
            row_cells[2].text = usage.replace('用法:', '').replace('用法：', '').strip()
        else:
            sig = get_function_signature(method)
            row_cells[2].text = f'''{name}{sig}'''
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, '仿宋', 10)
    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(10)
    doc.add_paragraph()


def add_function_details(doc, functions):
    '''添加函数详细说明'''
    for name, method in enumerate(functions, 1):
        (description, usage) = parse_docstring(method.__doc__)
        sig = get_function_signature(method)
        title_text = f'''（{idx}）{name}{sig}'''
        if description:
            title_text += f''' {description}'''
        para = doc.add_paragraph()
        run = para.add_run(title_text)
        set_run_font(run, '仿宋', 11, True)
        if usage:
            para = doc.add_paragraph()
            run = para.add_run(usage)
            set_run_font(run, '仿宋', 10.5)
        doc.add_paragraph()


def add_math_function_section(doc):
    '''添加数学函数部分'''
    add_title(doc, '一、数学函数', level = 1)
    para = doc.add_paragraph()
    run = para.add_run('数学函数用于基本的数学运算，包括三角函数、对数、取整、条件选择等操作。所有函数返回pd.Series类型。')
    set_run_font(run, '仿宋', 10.5)
    doc.add_paragraph()
    functions = get_functions_from_class(MathFunction)
    add_title(doc, '函数列表', level = 2)
    add_function_table(doc, functions)
    add_title(doc, '函数说明', level = 2)
    add_function_details(doc, functions)
    add_title(doc, 'API案例', level = 2)
    add_math_example(doc)


def add_math_example(doc):
    '''添加数学函数API案例'''
    code = "import AmazingData as ad\nimport config_user\nfrom AmazingData.operator.math_function import MathFunction\n\n# 登录\nad.login(username=config_user.user['username'],\n         password=config_user.user['password'],\n         host=config_user.user['host'],\n         port=config_user.user['port'])\n\n# 获取数据\nbase_data_object = ad.BaseData()\ncalendar = base_data_object.get_calendar()\nmarket_data_object = ad.MarketData(calendar)\ncode = '000001.SH'\nkline_day = market_data_object.query_kline([code], begin_date=20130101, end_date=20250722,\n                                           period=ad.constant.Period.day.value)\ndf = kline_day[code]\n\n# MAX - 求n个参数中的最大值\nresult = MathFunction.MAX(df['close'], df['open'])\n\n# MIN - 求n个参数中的最小值\nresult = MathFunction.MIN(df['close'], df['open'])\n\n# ACOS - 反余弦值\nresult = MathFunction.ACOS(df['close'] / df['close'].max())\n\n# ASIN - 反正弦值\nresult = MathFunction.ASIN(df['close'] / df['close'].max())\n\n# ATAN - 反正切值\nresult = MathFunction.ATAN(df['close'])\n\n# COS - 余弦值\nresult = MathFunction.COS(df['close'])\n\n# SIN - 正弦值\nresult = MathFunction.SIN(df['close'])\n\n# TAN - 正切值\nresult = MathFunction.TAN(df['close'])\n\n# EXP - e的X次幂\nresult = MathFunction.EXP(df['close'] / 1000)\n\n# LN - 自然对数\nresult = MathFunction.LN(df['close'])\n\n# LOG - 10为底的对数\nresult = MathFunction.LOG(df['close'])\n\n# SQRT - 开平方\nresult = MathFunction.SQRT(df['close'])\n\n# ABS - 绝对值\nresult = MathFunction.ABS(df['close'] - df['open'])\n\n# POW - 乘幂\nresult = MathFunction.POW(df['close'] / 1000, df['open'] / df['open'])\n\n# CEILING - 向上舍入\nresult = MathFunction.CEILING(df['close'])\n\n# FLOOR - 向下舍入\nresult = MathFunction.FLOOR(df['close'])\n\n# INTPART - 取整\nresult = MathFunction.INTPART(df['close'])\n\n# BETWEEN - 介于\nresult = MathFunction.BETWEEN(df['close'], df['low'], df['high'])\n\n# FRACPART - 小数部分\nresult = MathFunction.FRACPART(df['close'])\n\n# ROUND - 四舍五入\nresult = MathFunction.ROUND(df['close'], 2)\n\n# SIGN - 取符号\nresult = MathFunction.SIGN(df['close'] - df['open'])\n\n# MOD - 取模\nresult = MathFunction.MOD(df['close'], 100)\n\n# IF - 条件选择函数\nresult = MathFunction.IF(df['close'] > df['open'], df['high'], df['low'])\n\n# RAND - 取随机数\nresult = MathFunction.RAND(1, 100)"
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    doc.add_paragraph()


def add_statistics_function_section(doc):
    '''添加统计函数部分'''
    add_title(doc, '二、统计函数', level = 1)
    para = doc.add_paragraph()
    run = para.add_run('统计函数用于计算时序数据的统计指标，如标准差、方差、相关系数等。所有函数返回pd.Series类型。')
    set_run_font(run, '仿宋', 10.5)
    doc.add_paragraph()
    functions = get_functions_from_class(StatisticsFunction)
    add_title(doc, '函数列表', level = 2)
    add_function_table(doc, functions)
    add_title(doc, '函数说明', level = 2)
    add_function_details(doc, functions)
    add_title(doc, 'API案例', level = 2)
    add_statistics_example(doc)


def add_statistics_example(doc):
    '''添加统计函数API案例'''
    code = "import AmazingData as ad\nimport config_user\nfrom AmazingData.operator.statistics_function import StatisticsFunction\n\n# 登录并获取数据(同上)\n# ...\n# df = kline_day[code]\n\n# AVEDEV - 平均绝对偏差\nresult = StatisticsFunction.AVEDEV(df['close'], 20)\n\n# DEVSQ - 数据偏差平方和\nresult = StatisticsFunction.DEVSQ(df['close'], 20)\n\n# FORCAST - 线性回归预测值\nresult = StatisticsFunction.FORCAST(df['close'], 20)\n\n# SLOPE - 线性回归斜率\nresult = StatisticsFunction.SLOPE(df['close'], 20)\n\n# STD - 估算标准差(样本标准差)\nresult = StatisticsFunction.STD(df['close'], 20)\n\n# STDP - 总体标准差\nresult = StatisticsFunction.STDP(df['close'], 20)\n\n# STDDEV - 标准偏差\nresult = StatisticsFunction.STDDEV(df['close'], 20)\n\n# VAR - 估算样本方差\nresult = StatisticsFunction.VAR(df['close'], 20)\n\n# VARP - 总体样本方差\nresult = StatisticsFunction.VARP(df['close'], 20)\n\n# COVAR - 协方差\nresult = StatisticsFunction.COVAR(df['close'], df['open'], 20)\n\n# RELATE - 相关系数\nresult = StatisticsFunction.RELATE(df['close'], df['open'], 20)\n\n# BETA - 贝塔系数\nresult = StatisticsFunction.BETA(df['close'], df['open'], 20)\n\n# BETAEX - 相关放大系数\nresult = StatisticsFunction.BETAEX(df['close'], df['open'], 20)\n\n# KURTOSIS - 峰度\nresult = StatisticsFunction.KURTOSIS(df['close'], 20)\n\n# SKEW - 偏度\nresult = StatisticsFunction.SKEW(df['close'], 20)\n\n# MEAN - 平均值\nresult = StatisticsFunction.MEAN(df['close'], 20)\n\n# MEDIAN - 中位数\nresult = StatisticsFunction.MEDIAN(df['close'], 20)\n\n# QUANTILE - 分位数\nresult = StatisticsFunction.QUANTILE(df['close'], 20, 0.75)"
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    doc.add_paragraph()


def add_time_series_function_section(doc):
    '''添加时序函数部分'''
    add_title(doc, '三、时序函数', level = 1)
    para = doc.add_paragraph()
    run = para.add_run('时序函数用于时间序列数据的处理，包括引用、移动平均、条件统计、累计求和、技术指标等。')
    set_run_font(run, '仿宋', 10.5)
    doc.add_paragraph()
    functions = get_functions_from_class(TimeSeriesFunction)
    add_title(doc, '函数列表', level = 2)
    add_function_table(doc, functions)
    add_title(doc, '函数说明', level = 2)
    add_function_details(doc, functions)
    add_title(doc, 'API案例', level = 2)
    add_time_series_example(doc)


def add_time_series_example(doc):
    '''添加时序函数API案例'''
    code = "import AmazingData as ad\nimport config_user\nfrom AmazingData.operator.time_series_function import TimeSeriesFunction\n\n# 登录并获取数据(同上)\n# ...\n# df = kline_day[code]\n# condition = df['close'] > df['open']\n\n# ========== 位置信息函数 ==========\n# BARSTATUS - 返回数据的位置信息\nresult = TimeSeriesFunction.BARSTATUS(df['close'])\n\n# CURRBARSCOUNT - 求到最后K线的周期数\nresult = TimeSeriesFunction.CURRBARSCOUNT(df['close'])\n\n# TOTALBARSCOUNT - 求到当前的周期数(从1开始递增)\nresult = TimeSeriesFunction.TOTALBARSCOUNT(df['close'])\n\n# ========== 条件周期统计函数 ==========\n# BARSLAST - 上一次条件成立到当前的周期数\ncondition = df['close'] > df['open']\nresult = TimeSeriesFunction.BARSLAST(condition)\n\n# BARSLASTS - 倒数第N次成立时距今的周期数\nresult = TimeSeriesFunction.BARSLASTS(condition, 3)\n\n# BARSNEXT - 下一次条件成立到当前的周期数(未来函数)\nresult = TimeSeriesFunction.BARSNEXT(condition)\n\n# BARSSINCEN - N周期内第一个条件成立到当前的周期数\nresult = TimeSeriesFunction.BARSSINCEN(condition, 10)\n\n# BARSSINCE - 第一个条件成立到当前的周期数\nresult = TimeSeriesFunction.BARSSINCE(condition)\n\n# COUNT - 统计满足条件的周期数\nresult = TimeSeriesFunction.COUNT(condition, 20)\n\n# BARSLASTCOUNT - 统计连续满足条件的周期数\nresult = TimeSeriesFunction.BARSLASTCOUNT(condition)\n\n# ========== 最值函数 ==========\n# HHV - 求N周期内最高值\nresult = TimeSeriesFunction.HHV(df['high'], 20)\n\n# HHVBARS - 求上一高点到当前的周期数\nresult = TimeSeriesFunction.HHVBARS(df['high'], 20)\n\n# HOD - 求高值名次\nresult = TimeSeriesFunction.HOD(df['high'], 20)\n\n# LLV - 求N周期内最低值\nresult = TimeSeriesFunction.LLV(df['low'], 20)\n\n# LLVBARS - 求上一低点到当前的周期数\nresult = TimeSeriesFunction.LLVBARS(df['low'], 20)\n\n# LOD - 求低值名次\nresult = TimeSeriesFunction.LOD(df['low'], 20)\n\n# HHVLLV - 阶段最高最低值\nresult = TimeSeriesFunction.HHVLLV(df['high'], 0, 20, 5)\n\n# ========== 引用函数 ==========\n# REVERSE - 求相反数\nresult = TimeSeriesFunction.REVERSE(df['close'])\n\n# REF - 引用若干周期前的数据\nresult = TimeSeriesFunction.REF(df['close'], 1)\n\n# REFX - 引用若干周期后的数据(未来函数)\nresult = TimeSeriesFunction.REFX(df['close'], 1)\n\n# REFV - 引用若干周期前的数据(平滑处理)\nresult = TimeSeriesFunction.REFV(df['close'], 1)\n\n# REFXV - 引用若干周期后的数据(平滑处理)(未来函数)\nresult = TimeSeriesFunction.REFXV(df['close'], 1)\n\n# SHIFT - 获取N个交易日前的值\nresult = TimeSeriesFunction.SHIFT(df['close'], 5)\n\n# ========== 累计函数 ==========\n# SUM - 求总和\nresult = TimeSeriesFunction.SUM(df['volume'], 20)\n\n# MULAR - 求累乘\nresult = TimeSeriesFunction.MULAR(df['close'] / TimeSeriesFunction.REF(df['close'], 1), 5)\n\n# SUMBARS - 向前累加到指定值到现在的周期数\nresult = TimeSeriesFunction.SUMBARS(df['volume'], 1000000000)\n\n# SUMBARSX - 向前累加到指定值到现在的周期数(未达到返回nan)\nresult = TimeSeriesFunction.SUMBARSX(df['volume'], 1000000000)\n\n# CUMSUM - 累计求和\nresult = TimeSeriesFunction.CUMSUM(df['volume'])\n\n# ========== 移动平均函数 ==========\n# MA - 简单移动平均\nresult = TimeSeriesFunction.MA(df['close'], 20)\n\n# SMA - 移动平均\nresult = TimeSeriesFunction.SMA(df['close'], 20, 1)\n\n# TMA - 移动平均\nresult = TimeSeriesFunction.TMA(df['close'], 0.9, 0.1)\n\n# MEMA - 平滑移动平均\nresult = TimeSeriesFunction.MEMA(df['close'], 20)\n\n# EMA - 指数移动平均\nresult = TimeSeriesFunction.EMA(df['close'], 20)\n\n# EXPMEMA - 指数平滑移动平均\nresult = TimeSeriesFunction.EXPMEMA(df['close'], 20)\n\n# WMA - 加权移动平均\nresult = TimeSeriesFunction.WMA(df['close'], 20)\n\n# DMA - 动态移动平均\nalpha = df['volume'] / df['volume'].rolling(20).sum()\nresult = TimeSeriesFunction.DMA(df['close'], alpha)\n\n# AMA - 自适应均线值\nresult = TimeSeriesFunction.AMA(df['close'], alpha)\n\n# ========== 信号过滤函数 ==========\n# FILTER - 过滤连续出现的信号\nresult = TimeSeriesFunction.FILTER(condition, 5)\n\n# FILTERX - 反向过滤连续出现的信号\nresult = TimeSeriesFunction.FILTERX(condition, 5)\n\n# ========== 条件判断函数 ==========\n# TR - 求真实波幅\nresult = TimeSeriesFunction.TR(df['high'], df['low'], df['close'])\n\n# RANGE - 范围判断\nresult = TimeSeriesFunction.RANGE(df['close'], df['low'], df['high'])\n\n# CROSS - 两条线交叉\nma5 = TimeSeriesFunction.MA(df['close'], 5)\nma10 = TimeSeriesFunction.MA(df['close'], 10)\nresult = TimeSeriesFunction.CROSS(ma5, ma10)\n\n# LONGCROSS - 两条线维持一定周期后交叉\nresult = TimeSeriesFunction.LONGCROSS(ma5, ma10, 5)\n\n# UPNDAY - 返回周期数内是否连涨\nresult = TimeSeriesFunction.UPNDAY(df['close'], 3)\n\n# DOWNNDAY - 返回周期数内是否连跌\nresult = TimeSeriesFunction.DOWNNDAY(df['close'], 3)\n\n# NDAY - 返回是否持续存在X>Y\nresult = TimeSeriesFunction.NDAY(df['close'], df['open'], 3)\n\n# EXIST - 是否存在\nresult = TimeSeriesFunction.EXIST(condition, 10)\n\n# EXISTR - 是否存在(前几日到前几日间)\nresult = TimeSeriesFunction.EXISTR(condition, 10, 5)\n\n# EVERY - 一直存在\nresult = TimeSeriesFunction.EVERY(condition, 5)\n\n# LAST - 持续存在\nresult = TimeSeriesFunction.LAST(condition, 10, 5)\n\n# ========== 技术指标函数 ==========\n# SAR - 抛物线转向指标\nresult = TimeSeriesFunction.SAR(df['high'], df['low'], df['close'], n=4, step=0.02, max_af=0.2)"
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    doc.add_paragraph()


def add_cross_section_function_section(doc):
    '''添加截面函数部分'''
    add_title(doc, '四、截面函数', level = 1)
    para = doc.add_paragraph()
    run = para.add_run('截面函数用于计算同一交易日内多个标的之间的统计指标。输入数据为DataFrame，行为日期，列为标的代码。')
    set_run_font(run, '仿宋', 10.5)
    doc.add_paragraph()
    functions = get_functions_from_class(CrossSectionFunction)
    add_title(doc, '函数列表', level = 2)
    add_function_table(doc, functions)
    add_title(doc, '函数说明', level = 2)
    add_function_details(doc, functions)
    add_title(doc, 'API案例', level = 2)
    add_cross_section_example(doc)


def add_cross_section_example(doc):
    '''添加截面函数API案例'''
    code = "import AmazingData as ad\nimport config_user\nimport pandas as pd\nfrom AmazingData.operator.cross_section_function import CrossSectionFunction\n\n# 登录并获取数据\n# ...\n# 多只股票数据 (用于截面函数)\ncodes = ['000001.SZ', '000002.SZ', '000004.SZ', '000005.SZ', '000006.SZ']\nkline_multi = market_data_object.query_kline(codes, begin_date=20240101, end_date=20250101,\n                                             period=ad.constant.Period.day.value)\n# 构建截面数据 DataFrame (行:日期, 列:标的)\nclose_df = pd.DataFrame({c: kline_multi[c]['close'] for c in codes if c in kline_multi})\nopen_df = pd.DataFrame({c: kline_multi[c]['open'] for c in codes if c in kline_multi})\n\n# CSCOV - 截面协方差\nresult = CrossSectionFunction.CSCOV(close_df, open_df)\n\n# CSCOUNT - 截面标的个数\nresult = CrossSectionFunction.CSCOUNT(close_df)\n\n# CSQUANTILE - 截面分位数\nresult = CrossSectionFunction.CSQUANTILE(close_df, 0.5)\n\n# CSRANK - 截面排名\nresult = CrossSectionFunction.CSRANK(close_df, ascending=True)\n\n# CSSTD - 截面标准差\nresult = CrossSectionFunction.CSSTD(close_df)\n\n# CSSUM - 截面求和\nresult = CrossSectionFunction.CSSUM(close_df)\n\n# CSVAR - 截面方差\nresult = CrossSectionFunction.CSVAR(close_df)\n\n# CSPCTRANK - 截面百分位排名\nresult = CrossSectionFunction.CSPCTRANK(close_df)\n\n# CSMEAN - 截面平均值\nresult = CrossSectionFunction.CSMEAN(close_df)\n\n# CSMAX - 截面最大值\nresult = CrossSectionFunction.CSMAX(close_df)\n\n# CSCORR - 截面相关度\nresult = CrossSectionFunction.CSCORR(close_df, open_df)\n\n# CSMIN - 截面最小值\nresult = CrossSectionFunction.CSMIN(close_df)\n\n# CSMEDIAN - 截面中位数\nresult = CrossSectionFunction.CSMEDIAN(close_df)\n\n# CSZSCORE - 截面Z-score标准化\nresult = CrossSectionFunction.CSZSCORE(close_df)\n\n# CSNORMALIZE - 截面归一化(Min-Max)\nresult = CrossSectionFunction.CSNORMALIZE(close_df)\n\n# CSDEMEAN - 截面去均值\nresult = CrossSectionFunction.CSDEMEAN(close_df)"
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    doc.add_paragraph()


def generate_doc():
    '''生成完整的文档'''
    doc = create_document()
    add_title(doc, '算子函数说明文档', level = 0)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('版本: 1.1')
    set_run_font(run, '仿宋', 10.5)
    doc.add_paragraph()
    add_math_function_section(doc)
    doc.add_page_break()
    add_statistics_function_section(doc)
    doc.add_page_break()
    add_time_series_function_section(doc)
    doc.add_page_break()
    add_cross_section_function_section(doc)
    output_path = '算子函数说明文档.docx'
    doc.save(output_path)
    print(f'''文档已生成: {output_path}''')
    return output_path

if __name__ == '__main__':
    generate_doc()
    return None
